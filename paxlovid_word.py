from docx import Document
from docx.oxml.shared import OxmlElement,qn
from docx.shared import Cm
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START 
import datetime
import tkinter as tk
import os

def ADtoROC(AD): #西元轉民國
    try:
        AD=datetime.datetime.strptime(str(AD),"%Y/%m/%d")
    except:
        AD=datetime.datetime.strptime(str(AD),"%Y-%m-%d")
    ROC=str(AD.year-1911)+str(AD.strftime('%m'))+str(AD.strftime('%d'))
    if len(ROC)==6:
        ROC='0'+ROC
    return ROC

def ROC_today(): #民國今日日期
    return ADtoROC(datetime.date.today())

def save_word(ddi_result,med_table):
    def set_repeat_table_header(row): #重複標題列
        """ set repeat table row on every new page
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row
    docx = Document(r'files/輸出word_demo.docx') #引入範本
    docx.styles['Normal'].font.name = u'標楷體' #設定字體
    docx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') #設定字體及字形
    ddi = docx.add_paragraph(ddi_result)
    docx.add_paragraph() 
    # 添加一個空白段落
    section = docx.add_section(start_type=WD_SECTION_START.CONTINUOUS) 
    # 添加橫向頁的連續節
    section.orientation = WD_ORIENTATION.LANDSCAPE 
    # 設置橫向
    page_h, page_w = section.page_width, section.page_height
    section.page_width = page_w 
    # 設置橫向紙的寬度
    section.page_height = page_h 
    # 設置橫向紙的高度
    docx.add_paragraph('以下為病人用藥清單')
    t=docx.add_table(med_table.shape[0]+1,med_table.shape[1],style="Table Grid") #做一個和dataframe相同欄數和列數的表格，style="Table Grid"是全部都有格線
    for j in range(med_table.shape[-1]): #插入標題列文字
        t.cell(0,j).text=med_table.columns[j]
    for i in range(med_table.shape[0]): #插入內容
        for j in range(med_table.shape[-1]):
            t.cell(i+1,j).text=str(med_table.values[i,j])
    set_repeat_table_header(t.rows[0]) #重複宣告標題列
    
    #儲存檔案
    file_name=tk.filedialog.asksaveasfilename(title='另存word',initialfile=ROC_today(),defaultextension="*.docx", filetypes=[("DOCX", ".docx")])
    docx.save(file_name)
    os.system(file_name)